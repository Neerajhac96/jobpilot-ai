import json

from docx import Document

from app.ai.groq_client import GroqClient


def extract_docx_text(file_path: str) -> str:
    document = Document(file_path)

    sections = []

    # Normal paragraphs
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            sections.append(text)

    # Tables
    for table in document.tables:
        for row in table.rows:
            cells = []

            for cell in row.cells:
                text = cell.text.strip()

                if text:
                    cells.append(text)

            if cells:
                sections.append(" | ".join(cells))

    return "\n".join(sections)


class ResumeParserAgent:

    def __init__(self):
        self.ai = GroqClient()

    def parse(self, resume_text: str) -> dict:

        if not resume_text.strip():
            raise ValueError(
                "Resume contains no readable text."
            )

        prompt = f"""
Extract structured candidate information from the resume.

You MUST return exactly one valid JSON object.

Do not write:
- Markdown
- ```json
- explanations
- notes
- introductions

Return ONLY JSON.

Use exactly this structure:

{{
    "name": "",
    "email": "",
    "phone": "",
    "location": "",
    "education": "",
    "skills": [],
    "experience": "",
    "projects": [],
    "certifications": []
}}

Rules:

1. Extract only information explicitly present in the resume.
2. Never invent information.
3. Do not infer missing skills.
4. Preserve project names and facts.
5. Preserve education accurately.
6. Preserve certifications accurately.
7. Preserve experience accurately.
8. Missing values must be empty strings or empty arrays.
9. Keep skills as individual items.
10. Keep projects as individual items.

RESUME:

{resume_text}
"""

        response = self.ai.generate(
            prompt,
            json_mode=True
        )

        try:
            parsed = json.loads(response)

        except json.JSONDecodeError as exc:
            raise ValueError(
                f"Resume parser returned invalid JSON: {response}"
            ) from exc

        required_fields = [
            "name",
            "email",
            "phone",
            "location",
            "education",
            "skills",
            "experience",
            "projects",
            "certifications"
        ]

        for field in required_fields:
            if field not in parsed:
                raise ValueError(
                    f"Resume parser response is missing field: {field}"
                )

        return parsed
