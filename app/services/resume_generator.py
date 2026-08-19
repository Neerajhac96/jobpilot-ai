import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUTPUT_DIR = Path("data/resumes")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def generate_resume_docx(
    candidate,
    job,
    tailored_resume: dict
) -> str:

    document = Document()

    # Page setup
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    # Default font
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    # Name
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = name.add_run(candidate.name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Arial"

    # Contact line
    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_text = " | ".join(
        value
        for value in [
            candidate.email,
            candidate.phone,
            candidate.location
        ]
        if value
    )

    run = contact.add_run(contact_text)
    run.font.size = Pt(9)

    # Section helper
    def add_section_heading(title: str):
        paragraph = document.add_paragraph()

        paragraph.paragraph_format.space_before = Pt(7)
        paragraph.paragraph_format.space_after = Pt(2)

        run = paragraph.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)

        return paragraph

    # Professional Summary
    summary = tailored_resume.get(
        "professional_summary",
        ""
    )

    if summary:
        add_section_heading("Professional Summary")

        paragraph = document.add_paragraph(summary)
        paragraph.paragraph_format.space_after = Pt(3)

    # Skills
    skills = tailored_resume.get("skills", [])

    if skills:
        add_section_heading("Skills")

        paragraph = document.add_paragraph(
            " | ".join(skills)
        )

        paragraph.paragraph_format.space_after = Pt(3)

    # Experience
    experience = tailored_resume.get(
        "experience",
        []
    )

    if experience:
        add_section_heading("Experience")

        for item in experience:

            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            paragraph.paragraph_format.space_after = Pt(1)

            paragraph.add_run(str(item))

    # Projects
    projects = tailored_resume.get(
        "projects",
        []
    )

    if projects:
        add_section_heading("Projects")

        for item in projects:

            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            paragraph.paragraph_format.space_after = Pt(1)

            paragraph.add_run(str(item))

    # Education
    education = tailored_resume.get(
        "education",
        []
    )

    if education:
        add_section_heading("Education")

        for item in education:

            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            paragraph.paragraph_format.space_after = Pt(1)

            paragraph.add_run(str(item))

    # Certifications
    certifications = tailored_resume.get(
        "certifications",
        []
    )

    if certifications:
        add_section_heading("Certifications")

        for item in certifications:

            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            paragraph.paragraph_format.space_after = Pt(1)

            paragraph.add_run(str(item))

    filename = (
        f"{candidate.id}_{job.id}_tailored_resume.docx"
    )

    output_path = OUTPUT_DIR / filename

    document.save(output_path)

    return str(output_path)
